"""
Mart 2026.docx → backend/app/data/fjsti_staff_mart2026.json

Ustunlar: F.I.Sh, lavozim, tug'ilgan sana; bo'lim — jadval bo'lim sarlavhalari.
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

try:
    import docx
except ImportError:
    import subprocess

    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    import docx

DEFAULT_DOCX = Path(r"c:\Users\alocomputers\Desktop\Mart 2026.docx")
OUT = Path(__file__).resolve().parents[1] / "app" / "data" / "fjsti_staff_mart2026.json"


def clean_section(raw: str) -> str:
    s = " ".join(raw.split())
    s = re.sub(r"^\d+\s*ta\s*", "", s, flags=re.I)
    s = re.sub(r"\s+\d+\s*ta\s*", " ", s, flags=re.I)
    s = re.sub(r"\s+", " ", s).strip(" ·-–—")
    return s


def split_fio(full: str) -> tuple[str, str, str | None]:
    # "Salimov Nosir Inomjon o'g'li daktarant" → strip noise
    full = re.sub(r"\s+", " ", full).strip()
    full = re.sub(r"\s+(daktarant|doktorant|aspirant)\s*$", "", full, flags=re.I)
    parts = full.split()
    if not parts:
        return "", "", None
    if len(parts) == 1:
        return parts[0], "-", None
    if len(parts) == 2:
        return parts[0], parts[1], None
    # familiya + ism + otasining ismi (+ o'g'li/qizi)
    last = parts[0]
    first = parts[1]
    middle = " ".join(parts[2:]) or None
    return last, first, middle


def parse_birth(raw: str) -> str | None:
    if not raw:
        return None
    s = " ".join(raw.split())
    s = re.sub(r"[^\d./\-]", "", s)
    s = s.strip("./-")
    if not s:
        return None
    m = re.fullmatch(r"(\d{1,2})[./\-](\d{1,2})[./\-](\d{4})", s)
    if m:
        d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if 1 <= d <= 31 and 1 <= mo <= 12 and 1930 <= y <= 2015:
            return f"{y:04d}-{mo:02d}-{d:02d}"
        return None
    m = re.fullmatch(r"(\d{4})", s)
    if m:
        y = int(m.group(1))
        if 1930 <= y <= 2015:
            return f"{y:04d}-01-01"
    return None


def parse_docx(path: Path) -> list[dict]:
    doc = docx.Document(str(path))
    if not doc.tables:
        raise SystemExit("Jadval topilmadi")
    table = doc.tables[0]
    current: str | None = None
    by_dept: dict[str, list[dict]] = {}

    for row in table.rows[1:]:
        cells = [(" ".join(c.text.split())).strip() for c in row.cells]
        while len(cells) < 5:
            cells.append("")
        name, position, birth = cells[1], cells[3], cells[4]
        nonempty = [c for c in cells if c]
        if nonempty and len(set(nonempty)) == 1:
            current = clean_section(nonempty[0])
            by_dept.setdefault(current, [])
            continue
        if not name or name.lower() in ("f.i.sh", "№", "no"):
            continue
        last, first, middle = split_fio(name)
        if not last or not first:
            continue
        person = {
            "full_name_raw": name,
            "last_name": last,
            "first_name": first,
            "middle_name": middle,
            "position": position or None,
            "birth_date": parse_birth(birth),
        }
        key = current or "Noma'lum bo'lim"
        by_dept.setdefault(key, []).append(person)

    return [{"department": dept, "people": people} for dept, people in by_dept.items()]


def main() -> None:
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_DOCX
    if not src.exists():
        raise SystemExit(f"Fayl topilmadi: {src}")
    blocks = parse_docx(src)
    total = sum(len(b["people"]) for b in blocks)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    OUT.write_text(json.dumps(blocks, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"OK: {len(blocks)} bolim, {total} xodim -> {OUT}")


if __name__ == "__main__":
    main()
